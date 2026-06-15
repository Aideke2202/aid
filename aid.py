import streamlit as st
import os
import json
from google import genai
from docx import Document
from pypdf import PdfReader
import pandas as pd
import io
import sys

# Принудительно настраиваем окружение на UTF-8 для защиты от ошибок кодировки
import codecs
if sys.stdout.encoding != 'utf-8':
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
    sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')

# Настройка страницы
st.set_page_config(page_title="База знаний и Диспетчер ИИ", page_icon="🧠", layout="wide")

st.title("🧠 Интеллектуальная база знаний и Диспетчер СЗ")
st.write("Система анализа мануалов, таблиц Excel и автоматического распределения служебных записок по должностным инструкциям.")

# --- ГАРАНТИРОВАННОЕ СОЗДАНИЕ ПАПОК ---
BASE_DIR = "knowledge_base"
DOCS_DIR = os.path.join(BASE_DIR, "documents")
os.makedirs(DOCS_DIR, exist_ok=True)

# --- ИНИЦИАЛИЗАЦИЯ GEMINI (СТРОГО ВРУЧНУЮ) ---
st.sidebar.header("🔑 Настройки ИИ")

api_key = st.sidebar.text_input(
    "Введите Gemini API Key:", 
    type="password", 
    placeholder="Вставьте ключ AIzaSy..."
)

client = None
if api_key.strip():
    try:
        client = genai.Client(api_key=api_key.strip())
    except Exception as e:
        st.sidebar.error(f"Ошибка инициализации ИИ: {e}")
else:
    st.sidebar.warning("🔑 Пожалуйста, введите ваш рабочий API ключ слева для активации ИИ.")

# --- ФУНКЦИИ ДЛЯ ЧТЕНИЯ ФОРМАТОВ ---

def parse_docx(file_bytes):
    try:
        doc = Document(io.BytesIO(file_bytes))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    except Exception as e:
        # Если это старый .doc, замаскированный под .docx, пробуем вытащить из него текст напрямую
        try:
            text = file_bytes.decode('utf-8', errors='ignore')
            # Очищаем бинарный мусор старого Word
            clean_lines = [line.strip() for line in text.split('\n') if len(line.strip()) > 5]
            if len(clean_lines) > 2:
                return "\n".join(clean_lines)
        except:
            pass
        return f"Ошибка при чтении файла Word (.docx): {e}. Если у вас старый формат (.doc), пожалуйста, пересохраните его как .docx через обычный Word."

def parse_old_doc(file_bytes):
    """Резервный парсер для старых файлов .doc (вытаскивает строки текста напрямую из бинарника)"""
    try:
        # Извлекаем печатные символы из бинарного файла .doc
        text = ""
        for b in file_bytes:
            if 32 <= b <= 126 or 1040 <= b <= 1103 or b in [10, 13]: # ASCII + Кириллаца + перенос строки
                if b in [10, 13]:
                    text += "\n"
                else:
                    # Корректно декодируем байты кириллицы для старых кодировок Windows-1251
                    try:
                        text += bytes([b]).decode('cp1251', errors='ignore')
                    except:
                        text += chr(b)
        
        # Фильтруем системный мусор
        lines = [line.strip() for line in text.split('\n') if len(line.strip()) > 3]
        final_text = "\n".join(lines)
        
        if len(final_text.strip()) == 0:
            return "Ошибка: Не удалось извлечь текст из старого .doc файла. Пересохраните его в .docx"
        return final_text
    except Exception as e:
        return f"Ошибка при разборе старого Word (.doc): {e}"

def parse_pdf(file_bytes):
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        full_text = []
        for idx, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text:
                full_text.append(f"[Страница {idx}]\n{text}")
        return "\n".join(full_text)
    except Exception as e:
        return f"Ошибка при чтении файла PDF: {e}"

def parse_excel(file_bytes):
    try:
        excel_file = pd.ExcelFile(io.BytesIO(file_bytes), engine='openpyxl')
        full_text = []
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name=sheet_name, header=None, engine='openpyxl')
            df = df.fillna("")
            if not df.empty:
                full_text.append(f"--- Лист таблицы: {sheet_name} ---")
                for idx, row in df.iterrows():
                    row_values = []
                    for val in row.values:
                        val_str = str(val).strip()
                        if val_str in ["NaN", "NaT", "<NA>"]:
                            val_str = ""
                        row_values.append(val_str)
                    if any(row_values):
                        full_text.append(f"Строка Excel {idx+1}: " + " | ".join(row_values))
        return "\n".join(full_text)
    except Exception as e:
        return f"Ошибка при чтении файла Excel: {e}"

# --- БОКОВАЯ ПАНЕЛЬ: ЗАГРУЗКА И УДАЛЕНИЕ ---
with st.sidebar:
    st.write("---")
    st.header("📥 Загрузка инструкций и таблиц")
    
    doc_type = st.selectbox("Тип документа:", ["Должностная инструкция", "Инструкция к программе / Ошибка", "Таблица Excel / Реестр"])
    doc_title = st.text_input("Название (ФИО, программа или имя таблицы):")
    
    # ТЕПЕРЬ ОФИЦИАЛЬНО РАЗРЕШИЛИ РАСШИРЕНИЕ 'doc' В ИНТЕРФЕЙСЕ
    uploaded_file = st.file_uploader("Выложите файл (.docx, .doc, .pdf, .xlsx):", type=["docx", "doc", "pdf", "xlsx"])
    
    save_btn = st.button("💾 Сохранить в базу знаний", type="secondary")
    
    if save_btn and doc_title and uploaded_file:
        file_bytes = uploaded_file.read()
        file_ext = uploaded_file.name.split(".")[-1].lower()
        
        if file_ext == "docx":
            parsed_text = parse_docx(file_bytes)
        elif file_ext == "doc":
            # Активируем резервный бинарный парсер кириллицы для старых вордовских файлов
            parsed_text = parse_old_doc(file_bytes)
        elif file_ext == "pdf":
            parsed_text = parse_pdf(file_bytes)
        elif file_ext in ["xlsx", "xls"]:
            parsed_text = parse_excel(file_bytes)
        else:
            parsed_text = "Неподдерживаемый формат файла."
        
        if "Ошибка при чтении" in parsed_text or "Ошибка:" in parsed_text:
            st.error(parsed_text)
        else:
            clean_title = "".join(c for c in doc_title if c.isalnum() or c in "._- ")
            clean_type = doc_type.replace(' ', '_')
            
            filename = f"{clean_type}_{clean_title.replace(' ', '_')}.json"
            filepath = os.path.join(DOCS_DIR, filename)
            
            data = {
                "type": doc_type,
                "title": doc_title,
                "content": parsed_text
            }
            
            os.makedirs(DOCS_DIR, exist_ok=True)
            
            with open(filepath, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=4)
            st.success(f"Документ '{doc_title}' успешно добавлен!")
            st.rerun()

    # --- СЕКЦИЯ УДАЛЕНИЯ ДОКУМЕНТОВ ---
    st.write("---")
    st.subheader("📚 Список документов в базе:")
    if os.path.exists(DOCS_DIR):
        files = [f for f in os.listdir(DOCS_DIR) if f.endswith('.json')]
        if files:
            for file in files:
                filepath = os.path.join(DOCS_DIR, file)
                with open(filepath, "r", encoding="utf-8") as f:
                    meta = json.load(f)
                
                if meta["type"] == "Должностная инструкция":
                    icon = "👤"
                elif meta["type"] == "Таблица Excel / Реестр":
                    icon = "📊"
                else:
                    icon = "💻"
                
                col1, col2 = st.columns([4, 1])
                with col1:
                    st.write(f"{icon} **{meta['title']}**")
                with col2:
                    if st.button("🗑️", key=file, help=f"Удалить {meta['title']}"):
                        os.remove(filepath)
                        st.success("Удалено!")
                        st.rerun()
        else:
            st.info("База знаний пуста. Загрузите файлы.")

# --- ГЛАВНАЯ СТРАНИЦА: ПЕРЕКЛЮЧЕНИЕ ВКЛАДОК ---
tab1, tab2 = st.tabs(["📝 Общий поиск и Анализ", "🎯 Диспетчер Служебных Записок (СЗ)"])

# --- ВКЛАДКА 1: ОБЩИЙ ПОИСК ---
with tab1:
    st.subheader("Поиск в базе знаний и таблицах")
    user_query = st.text_area("Вставьте ваш вопрос по таблицам Excel или техническим мануалам:", height=120, key="query_main")
    
    if st.button("🔍 Запустить анализ", type="primary", key="btn_main"):
        if not api_key.strip() or client is None:
            st.error("Ошибка: Пожалуйста, сначала введите действующий API-ключ Gemini в левой панели!")
        elif not user_query.strip():
            st.warning("Пожалуйста, введите текст запроса.")
        else:
            all_docs_context = ""
            if os.path.exists(DOCS_DIR):
                files = [f for f in os.listdir(DOCS_DIR) if f.endswith('.json')]
                for file in files:
                    with open(os.path.join(DOCS_DIR, file), "r", encoding="utf-8") as f:
                        meta = json.load(f)
                        all_docs_context += f"--- НАЧАЛО ДОКУМЕНТА: {meta['title']} ({meta['type']}) ---\n{meta['content']}\n--- КОНЕЦ ДОКУМЕНТА ---\n\n"
            
            if not all_docs_context:
                st.error("Ошибка: База знаний пуста! Загрузите файлы слева.")
            else:
                with st.spinner("Анализирую базу данных..."):
                    try:
                        system_instruction = (
                            "You are an AI assistant for corporate knowledge base analysis. "
                            "Analyze the user's request using the provided context. Provide answers in Russian."
                        )
                        full_prompt = f"Context:\n{all_docs_context}\n\nUser Request:\n{user_query}"
                        
                        response = client.models.generate_content(
                            model='gemini-2.5-flash',
                            contents=full_prompt,
                            config={"system_instruction": system_instruction, "temperature": 0.1}
                        )
                        st.success("🤖 Ответ сформирован!")
                        st.markdown(response.text)
                    except Exception as e:
                        st.error(f"Ошибка при вызове Gemini: {e}")

# --- ВКЛАДКА 2: АВТОМАТИЧЕСКИЙ ДИСПЕТЧЕР СЗ ---
with tab2:
    st.subheader("🤖 Распределение СЗ по должностным инструкциям")
    st.write("Вставьте текст входящей служебной записки. ИИ сопоставит её суть с загруженными должностными инструкциями сотрудников и определит исполнителя.")
    
    sz_text = st.text_area("Текст служебной записки (СЗ):", height=200, placeholder="Пример: 'Необходимо провести дефектовку оборудования на узле связи и составить акт списания...'")
    
    if st.button("🎯 Определить ответственного сотрудника", type="primary", key="btn_sz"):
        if not api_key.strip() or client is None:
            st.error("Ошибка: Пожалуйста, сначала введите действующий API-ключ Gemini в левой панели!")
        elif not sz_text.strip():
            st.warning("Пожалуйста, вставьте текст служебной записки.")
        else:
            di_context = ""
            if os.path.exists(DOCS_DIR):
                files = [f for f in os.listdir(DOCS_DIR) if f.endswith('.json')]
                for file in files:
                    with open(os.path.join(DOCS_DIR, file), "r", encoding="utf-8") as f:
                        meta = json.load(f)
                        if meta["type"] == "Должностная инструкция":
                            di_context += f"--- ДОЛЖНОСТНАЯ ИНСТРУКЦИЯ СОТРУДНИКА: {meta['title']} ---\n{meta['content']}\n\n"
            
            if not di_context:
                st.error("Ошибка: В базе знаний нет ни одной Должностной инструкции! Пожалуйста, загрузите ДИ сотрудников в левой панели.")
            else:
                with st.spinner("Джамиля сопоставляет СЗ с должностными инструкциями..."):
                    try:
                        sz_system_instruction = (
                            "You are an automated corporate task dispatcher. "
                            "Your job is to read an incoming service note (СЗ) and compare it with the provided job descriptions (Должностные инструкции). "
                            "Determine which employee is responsible for the task described in the СЗ. "
                            "Provide your response in Russian. Structure it like this:\n"
                            "1. Clearly state the name/title of the responsible employee.\n"
                            "2. Quote the exact lines from their job description that justify your choice.\n"
                            "3. Give a brief, clear explanation why this task belongs to them.\n"
                            "Be objective and strict. If no matching employee is found, state that clearly."
                        )
                        
                        sz_prompt = f"""Вот список должностных инструкций сотрудников отдела:
{di_context}

---
ТЕКСТ ВХОДЯЩЕЙ СЛУЖЕБНОЙ ЗАПИСКИ (СЗ):
"{sz_text}"
---

На основе должностных инструкций определи, кому распределить эту служебную записку. Дай развернутый ответ на русском языке.
"""
                        sz_prompt_clean = str(sz_prompt).encode('utf-8', errors='ignore').decode('utf-8')
                        
                        response = client.models.generate_content(
                            model='gemini-2.5-flash',
                            contents=sz_prompt_clean,
                            config={"system_instruction": sz_system_instruction, "temperature": 0.1}
                        )
                        
                        st.success("🎯 Ответственный определен!")
                        st.subheader("📋 Резолюция ИИ:")
                        st.markdown(response.text)
                        
                    except Exception as e:
                        st.error(f"Ошибка распределения: {e}")
